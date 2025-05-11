import openai
client = openai.OpenAI(api_key=st.session_state.api_key)
import streamlit as st
import fitz  # PyMuPDF
import os
from docx import Document
from tempfile import NamedTemporaryFile

try:
    import openai
except ImportError:
    st.error("Le module openai est introuvable. Vérifie ton fichier requirements.txt ou installe-le.")

# Interface
st.set_page_config(page_title="Générateur de compte-rendus médicaux", page_icon="🩺")
st.title("🩺 Générateur de compte-rendus médicaux")

# API Key
openai_api_key = st.text_input("Clé API OpenAI", type="password")

# Upload du fichier PDF
uploaded_file = st.file_uploader("Sélectionnez une note médicale (PDF)", type=["pdf"])

# Fonction : lire texte depuis PDF
def extract_text_from_pdf(file) -> str:
    with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(file.read())
        tmp_path = tmp_file.name

    doc = fitz.open(tmp_path)
    texte = ""
    for page in doc:
        texte += page.get_text()
    os.remove(tmp_path)
    return texte

# Exemples témoins
EXEMPLES = [
    {
        "note": open("ex1_note.pdf", "rb").read(),
        "cr": Document("ex1_cr.docx").paragraphs
    },
    {
        "note": open("ex2_note.pdf", "rb").read(),
        "cr": Document("ex2_cr.docx").paragraphs
    }
]

def format_examples():
    exemples = []
    for ex in EXEMPLES:
        note_text = extract_text_from_pdf(file=ex["note"])
        cr_text = "\n".join([p.text for p in ex["cr"]])
        exemples.append((note_text, cr_text))
    return exemples

# Prompt complet
def build_prompt(exemples, new_note_text):
    prompt = "Voici des exemples de transformation de notes médicales en texte de compte-rendu :\n\n"
    for i, (note, cr) in enumerate(exemples):
        prompt += f"Exemple {i+1} – Note :\n{note}\nCompte-rendu :\n{cr}\n\n"
    prompt += "Voici une nouvelle note à transformer :\n"
    prompt += new_note_text
    prompt += "\n\nMerci de rédiger uniquement le texte principal du compte-rendu, sans en-tête, nom, date ni signature."
    return prompt

# Bouton de génération
if st.button("Générer le compte-rendu"):
    if not uploaded_file:
        st.error("Veuillez sélectionner un fichier PDF.")
    elif not openai_api_key:
        st.error("Veuillez renseigner votre clé API OpenAI.")
    else:
        with st.spinner("Génération en cours..."):
            try:
                note_text = extract_text_from_pdf(uploaded_file)
                exemples = format_examples()
                prompt = build_prompt(exemples, note_text)

                openai.api_key = openai_api_key
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3
                )
                result = response.choices[0].message["content"].strip()
                st.success("Compte-rendu généré :")
                st.text_area("Résultat", result, height=400)
            except Exception as e:
                st.error(f"Erreur lors de la génération : {str(e)}")

